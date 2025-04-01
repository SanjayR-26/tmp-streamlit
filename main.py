import pandas as pd
import os
import io
import fitz  # PyMuPDF
import numpy as np
from PIL import Image
from typing import List, Dict, Optional, Union, BinaryIO
import json
import easyocr
from fastapi import FastAPI, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
import openai
from docx import Document
from dotenv import load_dotenv
import openpyxl
from io import BytesIO
import logging
import time
import base64

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("data_capture_engine")

load_dotenv()

class DataCaptureEngine:
    def __init__(self, assets_csv_path: str, openai_api_key: str):
        """
        Initialize the DataCaptureEngine
        
        Args:
            assets_csv_path: Path to the CSV file containing asset types and fields
            openai_api_key: OpenAI API key
        """
        logger.info(f"Initializing DataCaptureEngine with assets CSV: {assets_csv_path}")
        self.assets_df = pd.read_csv(assets_csv_path, sep='\t')
        self.reader = easyocr.Reader(['en'])
        self.openai_api_key = openai_api_key
        openai.api_key = openai_api_key
        logger.info(f"Loaded {len(self.assets_df)} asset types from CSV")
        
    def get_asset_types(self) -> List[Dict]:
        """Return all available asset types from the CSV"""
        asset_types = []
        for _, row in self.assets_df.iterrows():
            asset_types.append({
                "broad_category": row["BroadCategory"],
                "category": row["Category"],
                "subcategory": row["Subcategory"],
                "type": row["Type"]
            })
        return asset_types
    
    def get_fields_for_asset_type(self, asset_type: str) -> List[str]:
        """Get the relevant fields for a specific asset type"""
        matching_row = self.assets_df[self.assets_df["Type"] == asset_type]
        if not matching_row.empty:
            # Convert the string representation of list to actual list
            fields_str = matching_row.iloc[0]["Relevant Fields"]
            # Remove brackets and quotes, split by comma
            fields = fields_str.strip("[]").replace("'", "").split(", ")
            return fields
        return []
    
    def document_to_text(self, file_content: bytes, filename: str) -> List[str]:
        """Extract text from various document types"""
        text_pages = []
        
        # Process based on file type
        if filename.lower().endswith('.pdf'):
            # PDF processing
            doc = fitz.open(stream=file_content, filetype="pdf")
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text_pages.append(page.get_text())
            doc.close()
            
        elif filename.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif')):
            # Image processing with OCR
            image = Image.open(BytesIO(file_content))
            # Convert to numpy array for easyocr
            image_np = np.array(image)
            result = self.reader.readtext(image_np)
            # Extract text from OCR results
            text = " ".join([detection[1] for detection in result])
            text_pages.append(text)
            
        elif filename.lower().endswith('.docx'):
            # DOCX processing
            doc = Document(BytesIO(file_content))
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(" | ".join(row_text))
            text_pages.append("\n".join(full_text))
            
        elif filename.lower().endswith(('.xlsx', '.xls')):
            # Excel processing
            workbook = openpyxl.load_workbook(BytesIO(file_content))
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    sheet_text.append(" | ".join(row_values))
                text_pages.append(f"Sheet: {sheet_name}\n" + "\n".join(sheet_text))
        
        return text_pages
    
    def pdf_to_images(self, pdf_bytes: bytes) -> List[Image.Image]:
        """Convert PDF pages to PIL images"""
        images = []
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        zoom = 2  # Higher zoom for better OCR
        mat = fitz.Matrix(zoom, zoom)
        
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            pixmap = page.get_pixmap(matrix=mat)
            img = Image.open(BytesIO(pixmap.tobytes()))
            images.append(img)
        
        doc.close()
        return images
    
    def process_document(self, file_content: bytes, filename: str, asset_type: str) -> Dict:
        """
        Process a document to extract information based on the asset type
        
        Args:
            file_content: Binary content of the file
            filename: Original filename with extension
            asset_type: The type of asset to extract information for
            
        Returns:
            Dictionary containing extracted fields
        """
        start_time = time.time()
        logger.info(f"Processing document: {filename} for asset type: {asset_type}")
        
        # Get relevant fields for the asset type
        fields = self.get_fields_for_asset_type(asset_type)
        if not fields:
            logger.error(f"No fields found for asset type: {asset_type}")
            return {"error": f"No fields found for asset type: {asset_type}"}
        
        logger.info(f"Extracted fields for asset type {asset_type}: {fields}")
        
        # Extract text from the document
        logger.info(f"Extracting text from {filename}")
        text_pages = self.document_to_text(file_content, filename)
        logger.info(f"Extracted {len(text_pages)} pages of text")
        
        # Convert documents to images
        images = []
        
        if filename.lower().endswith('.pdf'):
            # PDF to images
            logger.info("Converting PDF to images")
            images = self.pdf_to_images(file_content)
            logger.info(f"Converted PDF to {len(images)} images")
        
        elif filename.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif')):
            # Already an image
            logger.info("Document is already an image")
            image = Image.open(BytesIO(file_content))
            images = [image]
        
        elif filename.lower().endswith('.docx'):
            # DOCX to images (using docx2pdf and then pdf to images)
            logger.info("Converting DOCX to images")
            # For DOCX, we'll do text extraction only as direct conversion to images is complex
            # We could implement a docx2pdf converter here if needed
            # For now, creating a placeholder image with text
            for i, page_text in enumerate(text_pages):
                # Create a simple image with text for representation
                img = Image.new('RGB', (800, 1100), color=(255, 255, 255))
                from PIL import ImageDraw, ImageFont
                draw = ImageDraw.Draw(img)
                try:
                    font = ImageFont.truetype("arial.ttf", 12)
                except IOError:
                    font = ImageFont.load_default()
                
                # Draw text on image (simplified)
                draw.text((10, 10), f"Page {i+1} - {page_text[:500]}...", fill=(0, 0, 0), font=font)
                images.append(img)
        
        elif filename.lower().endswith(('.xlsx', '.xls')):
            # Excel to images
            logger.info("Converting Excel to images")
            # Similar approach to DOCX - creating placeholder images
            for i, page_text in enumerate(text_pages):
                img = Image.new('RGB', (800, 1100), color=(255, 255, 255))
                from PIL import ImageDraw, ImageFont
                draw = ImageDraw.Draw(img)
                try:
                    font = ImageFont.truetype("arial.ttf", 12)
                except IOError:
                    font = ImageFont.load_default()
                
                draw.text((10, 10), f"Sheet {i+1} - {page_text[:500]}...", fill=(0, 0, 0), font=font)
                images.append(img)
        
        # Perform OCR on images for better text extraction
        image_texts = []
        for i, img in enumerate(images):
            logger.info(f"Running OCR on image {i+1}")
            img_np = np.array(img)
            result = self.reader.readtext(img_np)
            image_text = " ".join([detection[1] for detection in result])
            image_texts.append(image_text)
        
        # Combine text from document parsing and OCR
        all_text = "\n\n".join(text_pages + image_texts)
        logger.info(f"Combined text length: {len(all_text)} characters")
        
        try:
            logger.info(f"All text: {all_text}")
            
            # Use GPT-4o with JSON response format and include images
            model = "gpt-4o"
            logger.info(f"Using model: {model} with JSON response format and images")
            
            # Create content array with both text and images
            user_content = [
                {"type": "text", "text": f"Extract the following information from this document for asset type: {asset_type}\n\nFields needed: {fields}\n\nDocument content:\n{all_text}\n\nRespond with a valid JSON object containing only the requested fields. If a field cannot be found, set its value to null. If multiple values found send as a list of dictionary as json object, if nothing send an empty json"}
            ]
            
            # Add up to 3 images to the content (to avoid making requests too large)
            if images:
                logger.info(f"Adding {min(3, len(images))} images to the request")
                for i, img in enumerate(images[:3]):  # Limit to first 3 images
                    buffered = BytesIO()
                    img.save(buffered, format="PNG")
                    image_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                    user_content.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_base64}"}})
            
            messages = [
                {"role": "system", "content": "You are a data extraction assistant. Extract the requested fields from the document and return them as a JSON object."},
                {"role": "user", "content": user_content}
            ]
            
            logger.info(f"Calling OpenAI API with model: {model}")
            response = openai.chat.completions.create(
                model=model,
                messages=messages,
                response_format={"type": "json_object"}
            )
            
            logger.info(f"Response: {response}")
            
            # With response_format=json_object, we can parse directly without handling markdown blocks
            content = response.choices[0].message.content
            result = json.loads(content)
            processing_time = time.time() - start_time
            logger.info(f"Result: {result}")
            logger.info(f"Successfully extracted data in {processing_time:.2f} seconds")
            return result
            
        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"Error extracting information after {processing_time:.2f} seconds: {str(e)}")
            return {"error": f"Error extracting information: {str(e)}"}

# Create FastAPI application
app = FastAPI(title="Asset Data Capture Engine", version="1.0.0")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize engine
@app.on_event("startup")
async def startup_event():
    logger.info("Starting Asset Data Capture Engine application")
    app.state.engine = DataCaptureEngine(
        assets_csv_path="assets.csv",
        openai_api_key=os.environ.get("OPENAI_API_KEY", "your-api-key-here")
    )
    logger.info("Application startup complete")

@app.get("/asset-types")
async def get_asset_types():
    """Get all available asset types"""
    logger.info("Request for asset types")
    types = app.state.engine.get_asset_types()
    logger.info(f"Returning {len(types)} asset types")
    return types

@app.post("/extract")
async def extract_data(
    file: UploadFile = File(...),
    asset_type: str = Form(...)
):
    """
    Extract information from a document based on the selected asset type
    
    Args:
        file: The uploaded document file (PDF, DOCX, Excel, or image)
        asset_type: The selected asset type from the CSV
    """
    logger.info(f"Extract request received: file={file.filename}, asset_type={asset_type}")
    file_content = await file.read()
    logger.info(f"File size: {len(file_content)} bytes")
    
    result = app.state.engine.process_document(
        file_content=file_content,
        filename=file.filename,
        asset_type=asset_type
    )
    
    if "error" in result:
        logger.error(f"Extraction failed: {result['error']}")
    else:
        logger.info(f"Extraction successful: {len(result)} fields extracted")
    
    return result

if __name__ == "__main__":
    import uvicorn
    logger.info("Starting server on http://0.0.0.0:8000")
    uvicorn.run(app, host="0.0.0.0", port=8000) 
